import io
import base64
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
import logging

app = Flask(__name__)
CORS(app) # Habilita CORS para todas las rutas

# Configuración básica de logging para ver los mensajes en la terminal
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def docx_to_html(doc):
    """
    Convierte un objeto de documento de Word a una cadena HTML simple.
    Ahora incluye la lógica para conservar el formato de negrita, cursiva y subrayado.
    """
    html_content_parts = []
    for paragraph in doc.paragraphs:
        html_paragraph = ""
        # Iterar sobre las "runs" (fragmentos de texto con el mismo formato)
        for run in paragraph.runs:
            text = run.text
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            html_paragraph += text
        
        # Si el párrafo contiene texto, añadirlo como una etiqueta <p>
        if html_paragraph.strip():
            html_content_parts.append(f"<p>{html_paragraph}</p>")
            
    # Si no hay contenido, proporcionar un mensaje por defecto
    if not html_content_parts:
        return "<p>El documento resultante está vacío o no se pudo extraer contenido para previsualizar.</p>"
        
    return "".join(html_content_parts)


@app.route('/process-document', methods=['POST'])
def process_document():
    logging.info("Solicitud recibida en /process-document")
    try:
        # 1. Verificar si se subió un archivo
        if 'file' not in request.files:
            logging.error("No se encontró el archivo en la solicitud.")
            return jsonify({'error': 'No se encontró el archivo en la solicitud.'}), 400

        file = request.files['file']
        format_text = request.form.get('format_text', '').strip()

        if file.filename == '':
            logging.error("No se seleccionó ningún archivo.")
            return jsonify({'error': 'No se seleccionó ningún archivo.'}), 400

        if not format_text:
            logging.error("El texto de formato a remover/cambiar está vacío.")
            return jsonify({'error': 'Por favor, especifica qué quieres remover o cambiar.'}), 400

        # 2. Leer el documento de Word desde la solicitud
        original_doc_bytes = io.BytesIO(file.read())
        document = Document(original_doc_bytes)
        logging.info(f"Documento '{file.filename}' cargado exitosamente.")

        # 3. Crear un nuevo documento para guardar el contenido modificado
        modified_document = Document()
        
        # 4. Procesar y copiar párrafos, excluyendo los que coinciden con el texto a remover
        for paragraph in document.paragraphs:
            if format_text.lower() not in paragraph.text.lower():
                # Copiar el párrafo y su formato al nuevo documento
                new_paragraph = modified_document.add_paragraph(paragraph.text, style=paragraph.style.name)
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.name:
                        new_run.font.name = run.font.name
                    if run.font.size:
                        new_run.font.size = run.font.size
            else:
                logging.info(f"Párrafo eliminado (contiene '{format_text}'): {paragraph.text[:50]}...")

        # 5. Guardar el documento modificado en memoria
        modified_doc_bytes = io.BytesIO()
        modified_document.save(modified_doc_bytes)
        modified_doc_bytes.seek(0) # Volver al inicio del stream
        logging.info("Documento modificado y guardado en memoria.")

        # 6. Codificar el documento modificado a Base64 para enviarlo al frontend
        encoded_docx = base64.b64encode(modified_doc_bytes.read()).decode('utf-8')

        # 7. Unir las partes HTML para la previsualización
        final_html_content = docx_to_html(modified_document)

        logging.info("Procesamiento completado exitosamente.")
        return jsonify({
            'message': 'Documento procesado exitosamente',
            'docx_base64': encoded_docx,
            'html_content': final_html_content
        }), 200

    except Exception as e:
        logging.exception("Error inesperado durante el procesamiento del documento.")
        return jsonify({'error': f'Ocurrió un error en el servidor: {str(e)}'}), 500

if __name__ == '__main__':
    # Ejecuta la aplicación Flask en el puerto 5000
    app.run(debug=True, port=5000)


